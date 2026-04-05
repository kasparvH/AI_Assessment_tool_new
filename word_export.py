"""
word_export.py — Genereert een professioneel Word-rapport (.docx) op basis van
de AI Readiness Assessment resultaten. Gebruikt python-docx.
"""
from datetime import date
from pathlib import Path
from io import BytesIO

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Straightable brand colours
NAVY   = RGBColor(0x1B, 0x3B, 0x6F)
ORANGE = RGBColor(0xF7, 0x94, 0x1D)
GREY   = RGBColor(0xF8, 0xF9, 0xFA)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x11, 0x11, 0x11)
RED    = RGBColor(0xDC, 0x35, 0x45)
GREEN  = RGBColor(0x28, 0xA7, 0x45)

TIER_LABELS = {
    "Emerging":     "Beginnend",
    "Developing":   "In ontwikkeling",
    "Accelerating": "Versnellend",
    "Leading":      "Leidend",
}

FRAMEWORK_LABELS = {
    "EU_AI_ACT":   "EU AI Act",
    "NIST_AI_RMF": "NIST AI RMF",
    "ISO_42001":   "ISO 42001",
    "AI_TRISM":    "AI TRiSM",
}


def _set_cell_bg(cell, hex_color: str):
    """Set background colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = NAVY
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p


def _add_body(doc: Document, text: str, bold: bool = False, italic: bool = False, colour: RGBColor = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_rule(doc: Document):
    """Horizontal orange rule via bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "F7941D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    return p


def _score_color(pct: float) -> str:
    if pct < 35:
        return "DC3545"
    if pct < 60:
        return "F7941D"
    if pct < 80:
        return "FFC107"
    return "28A745"


def _tier_nl(tier: str) -> str:
    return TIER_LABELS.get(tier, tier)


def _bar(pct: float, width: int = 20) -> str:
    filled = round(pct / 100 * width)
    return "█" * filled + "░" * (width - filled)


def generate_word_report(
    report_text: str,
    dim_summary: pd.DataFrame,
    overall: dict,
    compliance: dict,
    org_name: str,
    respondent_name: str,
    sector: str,
    logo_path: str,
    inconsistencies: list,
    output_path: str,
):
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Default font ──
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── Cover block ──
    if Path(logo_path).exists():
        doc.add_picture(logo_path, width=Cm(6))

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("AI Readiness Assessment Rapport")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    run2 = sub.add_run(org_name)
    run2.font.size = Pt(18)
    run2.font.color.rgb = ORANGE
    run2.font.bold = True

    meta = doc.add_paragraph()
    meta.add_run(f"Respondent: {respondent_name}   |   Sector: {sector}   |   Datum: {date.today().strftime('%d %B %Y')}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _add_rule(doc)
    doc.add_paragraph()

    # ── Overall score box (table 1×2) ──
    score_table = doc.add_table(rows=1, cols=2)
    score_table.style = "Table Grid"
    left = score_table.cell(0, 0)
    right = score_table.cell(0, 1)
    _set_cell_bg(left,  "1B3B6F")
    _set_cell_bg(right, "F7941D")

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(f"Totaalscore\n{overall['overall_score_0_5']:.1f} / 5.0  ({overall['overall_pct']:.0f}%)")
    lr.font.size = Pt(16)
    lr.font.bold = True
    lr.font.color.rgb = WHITE

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run(f"Maturiteitsniveau\n{_tier_nl(overall['maturity_tier'])}")
    rr.font.size = Pt(16)
    rr.font.bold = True
    rr.font.color.rgb = WHITE

    doc.add_paragraph()
    _add_rule(doc)

    # ── Dimension scores table ──
    _add_heading(doc, "Overzicht scores per dimensie", level=2)

    dim_table = doc.add_table(rows=1, cols=4)
    dim_table.style = "Table Grid"
    headers = ["Dimensie", "Score", "Percentage", "Niveau"]
    for i, h in enumerate(headers):
        cell = dim_table.rows[0].cells[i]
        _set_cell_bg(cell, "1B3B6F")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)

    for _, row in dim_summary.sort_values("score_pct", ascending=False).iterrows():
        cells = dim_table.add_row().cells
        cells[0].text = row["dimension"]
        cells[1].text = f"{row['score_0_5']:.1f} / 5.0"
        cells[2].text = f"{row['score_pct']:.0f}%"
        tier = _tier_label(row["score_pct"])
        cells[3].text = tier
        hex_col = _score_color(row["score_pct"])
        _set_cell_bg(cells[3], hex_col)
        for c in cells[3:]:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE
                    r.font.bold = True
        for c in cells:
            c.paragraphs[0].runs[0].font.size = Pt(10) if c.paragraphs[0].runs else None

    doc.add_paragraph()

    # ── Report narrative (from Claude) ──
    _add_rule(doc)
    _render_report_text(doc, report_text)

    # ── GAP analysis ──
    _add_rule(doc)
    _add_heading(doc, "GAP-analyse: Hiaten en ontwikkelpunten", level=2)
    _add_body(doc,
        "De GAP-analyse hieronder brengt in kaart welke dimensies het grootste verschil vertonen "
        "ten opzichte van het gewenste maturiteitsniveau (score ≥ 4.0 / 5.0, equivalent aan minimaal 80%). "
        "Dit zijn de gebieden die de meeste prioriteit verdienen voor gerichte interventies.",
        italic=True
    )
    doc.add_paragraph()

    sorted_dims = dim_summary.sort_values("score_pct").to_dict("records")
    gap_table = doc.add_table(rows=1, cols=5)
    gap_table.style = "Table Grid"
    gap_headers = ["Dimensie", "Huidige score", "Streefwaarde", "GAP", "Prioriteit"]
    for i, h in enumerate(gap_headers):
        cell = gap_table.rows[0].cells[i]
        _set_cell_bg(cell, "1B3B6F")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)

    for row in sorted_dims:
        current = row["score_0_5"]
        target  = 4.0
        gap     = round(target - current, 2)
        gap_pct = round(max(0, 80 - row["score_pct"]), 1)
        if gap_pct == 0:
            priority = "Op niveau"
            pri_col  = "28A745"
        elif gap_pct < 20:
            priority = "Laag"
            pri_col  = "FFC107"
        elif gap_pct < 40:
            priority = "Middel"
            pri_col  = "F7941D"
        else:
            priority = "Hoog"
            pri_col  = "DC3545"

        cells = gap_table.add_row().cells
        cells[0].text = row["dimension"]
        cells[1].text = f"{current:.1f} / 5.0  ({row['score_pct']:.0f}%)"
        cells[2].text = "4.0 / 5.0  (80%)"
        cells[3].text = f"-{gap:.1f}  (-{gap_pct:.0f}%)" if gap > 0 else "✓"
        cells[4].text = priority
        _set_cell_bg(cells[4], pri_col)
        for p in cells[4].paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.bold = True
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    # ── Document inconsistencies ──
    if inconsistencies:
        _add_rule(doc)
        _add_heading(doc, f"Documentbevindingen: {len(inconsistencies)} afwijking(en) geconstateerd", level=2)
        _add_body(doc,
            "Tijdens de analyse zijn de door de respondent geüploade documenten vergeleken met de gegeven antwoorden. "
            "De onderstaande afwijkingen verdienen nadere aandacht en verificatie voordat conclusies worden getrokken."
        )
        for i, inc in enumerate(inconsistencies, 1):
            sev = inc.get("severity", "medium")
            sev_nl = {"high": "Hoog", "medium": "Middel", "low": "Laag"}.get(sev, sev)
            col = {"high": RED, "medium": ORANGE, "low": RGBColor(0xFF, 0xC1, 0x07)}.get(sev, ORANGE)

            p = doc.add_paragraph()
            r = p.add_run(f"Afwijking {i} — Ernst: {sev_nl}")
            r.font.bold = True
            r.font.color.rgb = col
            r.font.size = Pt(11)

            _add_body(doc, f"Vraag: {inc.get('question_text', '')}", bold=True)
            _add_body(doc, f"Gegeven antwoord: {inc.get('selected_answer', '')}")
            _add_body(doc, f"Document stelt: \"{inc.get('document_evidence', '')}\"", italic=True)
            _add_body(doc, f"Toelichting: {inc.get('explanation', '')}")
            doc.add_paragraph()

    # ── Compliance table ──
    _add_rule(doc)
    _add_heading(doc, "Compliance & Governance dekking", level=2)
    _add_body(doc,
        "De onderstaande tabel toont in hoeverre de beantwoorde vragen aansluiten bij de vier "
        "toonaangevende AI-governancekaders. Een lage dekkingsgraad duidt op mogelijke blinde vlekken "
        "in de governance en compliance voorbereiding van de organisatie."
    )
    doc.add_paragraph()

    comp_table = doc.add_table(rows=1, cols=3)
    comp_table.style = "Table Grid"
    for i, h in enumerate(["Kader", "Vragen gedekt", "Gemiddelde score"]):
        cell = comp_table.rows[0].cells[i]
        _set_cell_bg(cell, "1B3B6F")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)

    for fw, data in compliance.items():
        cells = comp_table.add_row().cells
        cells[0].text = FRAMEWORK_LABELS.get(fw, fw)
        cells[1].text = f"{data['covered']} / {data['total']} ({round(data['covered']/data['total']*100) if data['total'] else 0}%)"
        cells[2].text = f"{data['avg_score']:.1f} / 5.0"
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    # ── Footer note ──
    _add_rule(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        f"Vertrouwelijk — Straightable Innovatie & Strategie — {date.today().strftime('%d %B %Y')}"
    )
    fr.font.size = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(output_path)


def _render_report_text(doc: Document, text: str):
    """Parse the Claude-generated markdown report into Word paragraphs."""
    for line in text.split("\n"):
        line = line.rstrip()
        if line.startswith("## "):
            _add_heading(doc, line[3:], level=2)
        elif line.startswith("# "):
            _add_heading(doc, line[2:], level=1)
        elif line.startswith("**") and line.endswith("**") and len(line) > 4:
            _add_body(doc, line[2:-2], bold=True)
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:])
            run.font.size = Pt(11)
        elif line.strip() == "":
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        else:
            # Inline bold: replace **text** within a line
            _add_mixed_paragraph(doc, line)


def _add_mixed_paragraph(doc: Document, line: str):
    """Handle lines with inline **bold** markers."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parts = line.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = Pt(11)
        run.font.bold = (i % 2 == 1)  # odd indices are between ** markers


def _tier_label(pct: float) -> str:
    if pct < 35:
        return "Beginnend"
    if pct < 60:
        return "In ontwikkeling"
    if pct < 80:
        return "Versnellend"
    return "Leidend"
